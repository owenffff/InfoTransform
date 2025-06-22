#!/usr/bin/env python3
"""
Test single DOCX file upload to verify if DOCX processing works
"""

import requests
from pathlib import Path

def test_single_docx():
    # Create a simple test DOCX file using python-docx
    try:
        from docx import Document
        
        # Create a simple DOCX file
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test document for InfoTransform.')
        doc.add_paragraph('It contains some sample text to process.')
        
        # Save it
        test_file = Path("test_single.docx")
        doc.save(test_file)
        print(f"Created test DOCX file: {test_file}")
        
        # Test with the single file endpoint
        url = "http://localhost:8000/api/transform"
        
        with open(test_file, 'rb') as f:
            files = {'file': ('test_single.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            data = {
                'model_key': 'document_metadata',
                'custom_instructions': ''
            }
            
            print("\nTesting single file endpoint...")
            response = requests.post(url, files=files, data=data)
            
            print(f"Status: {response.status_code}")
            if response.status_code == 200:
                result = response.json()
                print("Success! Result:")
                print(result)
            else:
                print(f"Error: {response.text}")
        
        # Clean up
        test_file.unlink()
        
    except ImportError:
        print("python-docx not installed. Creating a simple text file instead.")
        test_file = Path("test_single.txt")
        test_file.write_text("This is a test document for InfoTransform.")
        print(f"Created test text file: {test_file}")

if __name__ == "__main__":
    test_single_docx()
